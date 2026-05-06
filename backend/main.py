#!/usr/bin/env python3
"""
PPT Master — FastAPI Backend

Provides REST API for:
- File conversion (DOCX/PPTX/PDF → Markdown)
- Storyboard generation (Markdown → Word/PDF)
- Project management
- Settings management
- LLM integration (Gemma4 local + cloud providers)
"""

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from pathlib import Path
from datetime import datetime
import shutil
import uuid
import json
import os
import sys
import subprocess
import re

# Fix Windows console encoding
os.environ["PYTHONIOENCODING"] = "utf-8"

app = FastAPI(title="PPT Master API", version="1.0.0")

# CORS for Next.js frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://frontend:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Paths
BASE_DIR = Path(__file__).parent.parent
PROJECTS_DIR = BASE_DIR / "projects"
EXPORTS_DIR = BASE_DIR / "exports"
UPLOADS_DIR = BASE_DIR / "uploads"
SETTINGS_FILE = BASE_DIR / "settings.json"
MODELS_DIR = BASE_DIR / "models"

# Ensure directories exist
PROJECTS_DIR.mkdir(exist_ok=True)
EXPORTS_DIR.mkdir(exist_ok=True)
UPLOADS_DIR.mkdir(exist_ok=True)
MODELS_DIR.mkdir(exist_ok=True)

# Python executable (works on both Windows and Linux)
PYTHON = sys.executable


def load_settings() -> dict:
    """Load settings from JSON file."""
    if SETTINGS_FILE.exists():
        return json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
    return {
        "llm": {
            "provider": "ollama",
            "model": "gemma4:8b",
            "endpoint": "http://localhost:11434",
            "api_key": "",
        },
        "image_search": {
            "pexels_api_key": "",
            "pixabay_api_key": "",
        },
        "tts": {
            "provider": "edge-tts",
            "voice": "vi-VN-HoaiMyNeural",
        },
        "language": "vi",
    }


def save_settings(settings: dict):
    """Save settings to JSON file."""
    SETTINGS_FILE.write_text(json.dumps(settings, indent=2, ensure_ascii=False), encoding="utf-8")


def clean_markdown_text(text: str) -> str:
    """Clean markdown text for output — remove artifacts, normalize whitespace."""
    # Remove markdown image references
    text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Remove page numbers like "01 / 10" or "01/10"
    text = re.sub(r'\d{2}\s*/\s*\d{2}', '', text)
    # Remove github URLs
    text = re.sub(r'github\.com/\S+', '', text)
    # Remove "PPT Master - ..." footer patterns
    text = re.sub(r'PPT Master\s*-\s*.*', '', text)
    # Remove "Source: ..." lines
    text = re.sub(r'^- Source:.*$', '', text, flags=re.MULTILINE)
    # Remove "Total slides: ..." lines
    text = re.sub(r'^- Total slides:.*$', '', text, flags=re.MULTILINE)
    # Remove "Slide N" headers (from ppt_to_md.py)
    text = re.sub(r'^## Slide \d+.*$', '', text, flags=re.MULTILINE)
    # Remove "#### N.NN Speaker Notes" sections and their content
    text = re.sub(r'####\s+\d+\.\d+\s+Speaker Notes.*?(?=####|\n## |\Z)', '', text, flags=re.DOTALL | re.MULTILINE)
    # Remove standalone "Speaker Notes" lines
    text = re.sub(r'^.*Speaker Notes.*$', '', text, flags=re.MULTILINE)
    # Remove slide number patterns like "01 / 10" at end of lines
    text = re.sub(r'\s*\d{2}\s*/\s*\d{2}\s*$', '', text, flags=re.MULTILINE)
    # Remove "PPT Master - XXX" lines
    text = re.sub(r'^.*PPT Master\s*-\s*.*$', '', text, flags=re.MULTILINE)
    # Remove empty markdown headers
    text = re.sub(r'^#{1,4}\s*$', '', text, flags=re.MULTILINE)
    # Remove multiple blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Remove trailing whitespace
    text = '\n'.join(line.rstrip() for line in text.split('\n'))
    # Clean up remaining artifacts
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    return text.strip()


# ─────────────────────────────────────────────
# Health check
# ─────────────────────────────────────────────

@app.get("/api/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}


# ─────────────────────────────────────────────
# Settings
# ─────────────────────────────────────────────

@app.get("/api/settings")
async def get_settings():
    return load_settings()


@app.post("/api/settings")
async def update_settings(settings: dict):
    save_settings(settings)
    return {"status": "ok"}


# ─────────────────────────────────────────────
# Models
# ─────────────────────────────────────────────

@app.get("/api/models")
async def list_models():
    """List available local models."""
    models = []
    if MODELS_DIR.exists():
        for f in MODELS_DIR.iterdir():
            if f.suffix in ('.gguf', '.bin', '.safetensors'):
                models.append({
                    "name": f.name,
                    "size_mb": round(f.stat().st_size / 1024 / 1024, 1),
                    "path": str(f),
                })
    
    # Check Ollama models
    ollama_models = []
    try:
        result = subprocess.run(
            ["ollama", "list"], capture_output=True, text=True, timeout=5
        )
        if result.returncode == 0:
            for line in result.stdout.strip().split('\n')[1:]:
                parts = line.split()
                if parts:
                    ollama_models.append(parts[0])
    except Exception:
        pass
    
    return {
        "local_files": models,
        "ollama_models": ollama_models,
        "models_dir": str(MODELS_DIR),
    }


# ─────────────────────────────────────────────
# File Upload & Conversion
# ─────────────────────────────────────────────

CONVERTER_MAP = {
    ".pdf": "pdf_to_md.py",
    ".docx": "doc_to_md.py",
    ".doc": "doc_to_md.py",
    ".pptx": "ppt_to_md.py",
    ".xlsx": "excel_to_md.py",
    ".xlsm": "excel_to_md.py",
    ".html": "doc_to_md.py",
    ".htm": "doc_to_md.py",
    ".epub": "doc_to_md.py",
}


def run_converter(upload_path: Path, suffix: str) -> Path:
    """Run the appropriate converter script and return the output markdown path."""
    scripts_dir = BASE_DIR / "skills" / "ppt-master" / "scripts" / "source_to_md"
    
    converter_script = CONVERTER_MAP.get(suffix)
    if not converter_script:
        raise HTTPException(400, f"Unsupported format: {suffix}")
    
    script_path = scripts_dir / converter_script
    if not script_path.exists():
        raise HTTPException(500, f"Converter script not found: {script_path}")
    
    result = subprocess.run(
        [PYTHON, str(script_path), str(upload_path)],
        capture_output=True, text=True, cwd=str(BASE_DIR),
    )
    
    if result.returncode != 0:
        raise HTTPException(500, f"Conversion failed: {result.stderr[:500]}")
    
    # Find output markdown file
    md_path = upload_path.with_suffix(".md")
    if not md_path.exists():
        md_files = list(upload_path.parent.glob(f"{upload_path.stem}*.md"))
        if md_files:
            md_path = md_files[0]
        else:
            raise HTTPException(500, "Conversion produced no output file")
    
    return md_path


@app.post("/api/convert")
async def convert_file(
    file: UploadFile = File(...),
    output_format: str = Form("markdown"),
):
    """Convert uploaded file (DOCX/PPTX/PDF) to Markdown."""
    file_id = str(uuid.uuid4())[:8]
    upload_path = UPLOADS_DIR / f"{file_id}_{file.filename}"
    
    with open(upload_path, "wb") as f:
        content = await file.read()
        f.write(content)
    
    suffix = upload_path.suffix.lower()
    md_path = run_converter(upload_path, suffix)
    
    return {
        "status": "ok",
        "file_id": file_id,
        "filename": file.filename,
        "markdown_path": str(md_path),
        "markdown_preview": md_path.read_text(encoding="utf-8")[:2000],
    }


# ─────────────────────────────────────────────
# Storyboard Generation
# ─────────────────────────────────────────────

@app.post("/api/storyboard/generate")
async def generate_storyboard(
    file: UploadFile = File(...),
    language: str = Form("vi"),
    output_format: str = Form("docx"),
):
    """
    Generate a structured storyboard from a document.
    
    Steps:
    1. Convert file to Markdown
    2. Send to LLM for structured extraction
    3. Generate Word/PDF output with professional formatting
    """
    try:
        # Step 1: Convert to Markdown
        file_id = str(uuid.uuid4())[:8]
        original_name = Path(file.filename).stem
        upload_path = UPLOADS_DIR / f"{file_id}_{file.filename}"
        
        with open(upload_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        suffix = upload_path.suffix.lower()
        md_path = run_converter(upload_path, suffix)
        markdown_content = md_path.read_text(encoding="utf-8")
        
        # Clean the markdown content
        markdown_content = clean_markdown_text(markdown_content)
        
        # Step 2: Send to LLM for structured extraction
        settings = load_settings()
        llm_config = settings.get("llm", {})
        
        storyboard_md = await _extract_storyboard_via_llm(
            markdown_content, language, llm_config
        )
        
        # Clean the storyboard output
        storyboard_md = clean_markdown_text(storyboard_md)
        
        # Step 3: Generate output file with correct naming
        # Format: {input_name}_{DD-MM-YYYY}.{ext}
        date_str = datetime.now().strftime("%d-%m-%Y")
        output_filename = f"{original_name}_{date_str}"
        output_dir = EXPORTS_DIR / "storyboards"
        output_dir.mkdir(exist_ok=True)
        
        if output_format == "docx":
            output_path = output_dir / f"{output_filename}.docx"
            _generate_docx(storyboard_md, output_path, language, original_name)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            output_path = output_dir / f"{output_filename}.pdf"
            _generate_pdf(storyboard_md, output_path, language, original_name)
            media_type = "application/pdf"
        
        # Save to projects history
        project_dir = PROJECTS_DIR / original_name
        project_dir.mkdir(exist_ok=True)
        (project_dir / "storyboard.md").write_text(storyboard_md, encoding="utf-8")
        (project_dir / "source.txt").write_text(file.filename, encoding="utf-8")
        (project_dir / "output_format.txt").write_text(output_format, encoding="utf-8")
        (project_dir / "language.txt").write_text(language, encoding="utf-8")
        (project_dir / "created_at.txt").write_text(datetime.now().isoformat(), encoding="utf-8")
        # Copy output file to project
        import shutil
        shutil.copy2(str(output_path), str(project_dir / output_path.name))
        
        return FileResponse(
            path=str(output_path),
            media_type=media_type,
            filename=output_path.name,
        )
    
    except HTTPException:
        raise
    except Exception as e:
        return JSONResponse(
            status_code=500,
            content={"status": "error", "detail": str(e)},
        )


async def _extract_storyboard_via_llm(
    markdown_content: str, language: str, llm_config: dict
) -> str:
    """Send markdown to LLM and get structured storyboard."""
    import httpx
    
    lang_names = {"vi": "Vietnamese", "en": "English", "zh": "Chinese"}
    lang_name = lang_names.get(language, "Vietnamese")
    
    prompt = f"""You are a document analysis expert. Read the following document and extract key information into a structured storyboard format.

RULES:
1. Output language: {lang_name}
2. Use numbered headers: 1.0, 1.1, 1.2, 2.0, 2.1, etc.
3. Each section must have clear fields with values
4. Use tables for data comparison when appropriate
5. Keep it concise but comprehensive
6. Format: Markdown
7. DO NOT include: image references, URLs, page numbers, slide numbers, "Source:" lines, "Total slides:" lines
8. DO NOT include: Speaker Notes sections — these are presentation metadata, not content
9. Focus on CONTENT only — extract key facts, data, and insights
10. Remove any "Slide N" headers — they are presentation artifacts
11. Clean up any garbled or meaningless text

OUTPUT FORMAT:
# Storyboard: [Document Title]

## 1.0 [Main Topic 1]
### 1.1 [Sub-topic]
- **Field**: Value
- **Field**: Value

### 1.2 [Sub-topic]
| Column 1 | Column 2 | Column 3 |
|---|---|---|
| Data | Data | Data |

## 2.0 [Main Topic 2]
...

DOCUMENT CONTENT:
{markdown_content[:8000]}

Generate the structured storyboard now:"""

    provider = llm_config.get("provider", "ollama")
    
    if provider == "ollama":
        endpoint = llm_config.get("endpoint", "http://localhost:11434")
        model = llm_config.get("model", "gemma4:8b")
        
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f"{endpoint}/api/generate",
                    json={
                        "model": model,
                        "prompt": prompt,
                        "stream": False,
                    },
                )
                response.raise_for_status()
                return response.json().get("response", "")
        except Exception:
            return _fallback_storyboard(markdown_content, language)
    
    elif provider == "openai":
        api_key = llm_config.get("api_key", "")
        endpoint = llm_config.get("endpoint", "https://api.openai.com/v1")
        model = llm_config.get("model", "gpt-4")
        
        try:
            async with httpx.AsyncClient(timeout=120.0) as client:
                response = await client.post(
                    f"{endpoint}/chat/completions",
                    headers={"Authorization": f"Bearer {api_key}"},
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "max_tokens": 4000,
                    },
                )
                response.raise_for_status()
                return response.json()["choices"][0]["message"]["content"]
        except Exception:
            return _fallback_storyboard(markdown_content, language)
    
    else:
        return _fallback_storyboard(markdown_content, language)


def _fallback_storyboard(markdown_content: str, language: str) -> str:
    """Fallback: create basic storyboard from markdown headers."""
    lines = markdown_content.split("\n")
    sections = []
    current_section = []
    section_num = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        if line.startswith("# "):
            if current_section:
                sections.append("\n".join(current_section))
            section_num += 1
            current_section = [f"## {section_num}.0 {line[2:]}"]
        elif line.startswith("## "):
            sub_num = len([s for s in current_section if s.startswith("###")]) + 1
            current_section.append(f"### {section_num}.{sub_num} {line[3:]}")
        elif line.startswith("### "):
            sub_num = len([s for s in current_section if s.startswith("###")]) + 1
            current_section.append(f"#### {section_num}.{sub_num} {line[4:]}")
        elif line.startswith("- ") or line.startswith("* "):
            current_section.append(line)
        elif line.startswith("|"):
            current_section.append(line)
        elif len(line) > 5:  # Skip very short artifacts
            current_section.append(f"- {line}")
    
    if current_section:
        sections.append("\n".join(current_section))
    
    return "\n\n".join(sections)


def _generate_docx(content: str, output_path: Path, language: str, title: str = "Storyboard"):
    """Generate professional Word document from storyboard markdown.
    
    Format: A4, Times New Roman, proper tables, headers, centered layout.
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    
    doc = Document()
    
    # ── Page Setup: A4, margins ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4 width
    section.page_height = Cm(29.7)  # A4 height
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    
    # ── Default font: Times New Roman ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15
    
    # Set East Asian font
    rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
    if rFonts is None:
        rPr = style.element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Times New Roman"/>')
        rPr.append(rFonts)
    
    # ── Title ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_after = Pt(4)
    title_run = title_para.add_run(f'STORYBOARD')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    title_run.font.name = 'Times New Roman'
    
    # ── Subtitle ──
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.space_after = Pt(6)
    subtitle_run = subtitle_para.add_run(title)
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x5b, 0x8d, 0xef)
    subtitle_run.font.name = 'Times New Roman'
    
    # ── Date line ──
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.space_after = Pt(12)
    date_run = date_para.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    date_run.font.name = 'Times New Roman'
    
    # ── Horizontal line ──
    line_para = doc.add_paragraph()
    line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_run = line_para.add_run('─' * 60)
    line_run.font.size = Pt(8)
    line_run.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)
    
    # ── Parse and render content ──
    lines = content.split("\n")
    table_rows = []
    in_table = False
    
    for line in lines:
        line = line.strip()
        if not line:
            # Flush table if we have rows
            if table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
            continue
        
        # Table detection
        if line.startswith("|") and "---" in line:
            in_table = True
            continue
        elif line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                table_rows.append(cells)
            continue
        else:
            # Flush table if we have rows
            if table_rows:
                _add_table(doc, table_rows)
                table_rows = []
                in_table = False
        
        # Headers
        if line.startswith("# ") and not line.startswith("## "):
            h = doc.add_heading(line[2:], level=0)
            _style_heading(h, Pt(18), RGBColor(0x1a, 0x1a, 0x2e))
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=1)
            _style_heading(h, Pt(15), RGBColor(0x2c, 0x3e, 0x50))
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=2)
            _style_heading(h, Pt(13), RGBColor(0x34, 0x49, 0x5e))
        elif line.startswith("#### "):
            h = doc.add_heading(line[5:], level=3)
            _style_heading(h, Pt(12), RGBColor(0x5b, 0x8d, 0xef))
        elif line.startswith("- **") or line.startswith("* **"):
            # Field: Value format
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(3)
            text = line[2:] if line.startswith("- ") else line[2:]
            parts = text.split("**")
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                if i % 2 == 1:  # Bold parts
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    
    # Flush remaining table
    if table_rows:
        _add_table(doc, table_rows)
    
    # ── Footer ──
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.space_before = Pt(20)
    footer_run = footer_para.add_run('Generated by PPT Master — AI Storyboard Generator')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)
    footer_run.font.name = 'Times New Roman'
    
    doc.save(str(output_path))


def _style_heading(heading, size, color):
    """Apply consistent heading style."""
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = size
        run.font.color.rgb = color


def _add_table(doc, rows):
    """Add a professionally styled table to the document."""
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Style the table
    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                
                if i == 0:  # Header row
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="2C3E50"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                else:
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                    if i % 2 == 0:  # Alternating row colors
                        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F9FA"/>')
                        cell._tc.get_or_add_tcPr().append(shading)
    
    # Add spacing after table
    doc.add_paragraph()


def _generate_pdf(content: str, output_path: Path, language: str, title: str = "Storyboard"):
    """Generate professional PDF from storyboard markdown.
    
    Format: A4, Times New Roman, proper tables, headers.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        HRFlowable
    )
    from reportlab.lib.units import cm, mm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        topMargin=2.54*cm,
        bottomMargin=2.54*cm,
        leftMargin=2.54*cm,
        rightMargin=2.54*cm,
    )
    
    # Custom styles
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Title'],
        fontName='Times-Roman',
        fontSize=22,
        textColor=HexColor('#1a1a2e'),
        alignment=TA_CENTER,
        spaceAfter=4,
    )
    
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=14,
        textColor=HexColor('#5b8def'),
        alignment=TA_CENTER,
        spaceAfter=6,
    )
    
    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=10,
        textColor=HexColor('#888888'),
        alignment=TA_CENTER,
        spaceAfter=12,
    )
    
    h1_style = ParagraphStyle(
        'H1', parent=styles['Heading1'],
        fontName='Times-Bold', fontSize=15,
        textColor=HexColor('#2c3e50'),
        spaceBefore=16, spaceAfter=8,
    )
    
    h2_style = ParagraphStyle(
        'H2', parent=styles['Heading2'],
        fontName='Times-Bold', fontSize=13,
        textColor=HexColor('#34495e'),
        spaceBefore=12, spaceAfter=6,
    )
    
    h3_style = ParagraphStyle(
        'H3', parent=styles['Heading3'],
        fontName='Times-Bold', fontSize=12,
        textColor=HexColor('#5b8def'),
        spaceBefore=8, spaceAfter=4,
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=11,
        textColor=HexColor('#333333'),
        spaceBefore=2, spaceAfter=4,
        leading=14,
    )
    
    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=1*cm,
        bulletIndent=0.5*cm,
    )
    
    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontName='Times-Roman',
        fontSize=8,
        textColor=HexColor('#aaaaaa'),
        alignment=TA_CENTER,
        spaceBefore=20,
    )
    
    story = []
    
    # Title
    story.append(Paragraph('STORYBOARD', title_style))
    story.append(Paragraph(title, subtitle_style))
    story.append(Paragraph(
        f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}', date_style
    ))
    story.append(HRFlowable(
        width="100%", thickness=1, color=HexColor('#cccccc'),
        spaceAfter=12
    ))
    
    # Parse content
    lines = content.split("\n")
    table_data = []
    
    for line in lines:
        line = line.strip()
        if not line:
            if table_data:
                _add_pdf_table(story, table_data)
                table_data = []
            story.append(Spacer(1, 0.2*cm))
            continue
        
        # Table detection
        if line.startswith("|") and "---" in line:
            continue
        elif line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if cells:
                table_data.append(cells)
            continue
        else:
            if table_data:
                _add_pdf_table(story, table_data)
                table_data = []
        
        # Escape XML
        safe = line.replace("&", "&").replace("<", "<").replace(">", ">")
        
        # Convert **bold** to <b>bold</b>
        safe = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', safe)
        
        if line.startswith("# ") and not line.startswith("## "):
            story.append(Paragraph(safe[2:], h1_style))
        elif line.startswith("## "):
            story.append(Paragraph(safe[3:], h1_style))
        elif line.startswith("### "):
            story.append(Paragraph(safe[4:], h2_style))
        elif line.startswith("#### "):
            story.append(Paragraph(safe[5:], h3_style))
        elif line.startswith("- ") or line.startswith("* "):
            story.append(Paragraph(f"• {safe[2:]}", bullet_style))
        else:
            story.append(Paragraph(safe, body_style))
    
    if table_data:
        _add_pdf_table(story, table_data)
    
    # Footer
    story.append(HRFlowable(
        width="100%", thickness=0.5, color=HexColor('#cccccc'),
        spaceBefore=20, spaceAfter=8
    ))
    story.append(Paragraph(
        'Generated by PPT Master — AI Storyboard Generator', footer_style
    ))
    
    doc.build(story)


def _add_pdf_table(story, rows):
    """Add a professionally styled table to PDF."""
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import Table, TableStyle, Spacer
    from reportlab.lib.units import cm
    
    if not rows:
        return
    
    # Create table
    table = Table(rows)
    
    # Style
    style_commands = [
        ('FONTNAME', (0, 0), (-1, -1), 'Times-Roman'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (-1, -1), HexColor('#333333')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#dddddd')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ('LEFTPADDING', (0, 0), (-1, -1), 6),
        ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        # Header row
        ('BACKGROUND', (0, 0), (-1, 0), HexColor('#2c3e50')),
        ('TEXTCOLOR', (0, 0), (-1, 0), HexColor('#ffffff')),
        ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
    ]
    
    # Alternating row colors
    for i in range(1, len(rows)):
        if i % 2 == 0:
            style_commands.append(('BACKGROUND', (0, i), (-1, i), HexColor('#f8f9fa')))
    
    table.setStyle(TableStyle(style_commands))
    story.append(table)
    story.append(Spacer(1, 0.3*cm))


# ─────────────────────────────────────────────
# Projects
# ─────────────────────────────────────────────

@app.get("/api/projects")
async def list_projects():
    """List all projects."""
    projects = []
    if PROJECTS_DIR.exists():
        for item in sorted(PROJECTS_DIR.iterdir(), reverse=True):
            if item.is_dir():
                # Check for storyboard projects
                has_storyboard = (item / "storyboard.md").exists()
                has_svg = (item / "svg_output").exists()
                has_spec = (item / "design_spec.md").exists()
                
                # Get metadata
                source_file = ""
                created_at = ""
                output_format = ""
                language = ""
                
                if (item / "source.txt").exists():
                    source_file = (item / "source.txt").read_text(encoding="utf-8").strip()
                if (item / "created_at.txt").exists():
                    created_at = (item / "created_at.txt").read_text(encoding="utf-8").strip()
                if (item / "output_format.txt").exists():
                    output_format = (item / "output_format.txt").read_text(encoding="utf-8").strip()
                if (item / "language.txt").exists():
                    language = (item / "language.txt").read_text(encoding="utf-8").strip()
                
                # Find output files
                output_files = []
                for f in item.iterdir():
                    if f.suffix in ('.docx', '.pdf', '.pptx'):
                        output_files.append({
                            "name": f.name,
                            "size": f.stat().st_size,
                        })
                
                projects.append({
                    "name": item.name,
                    "path": str(item),
                    "has_storyboard": has_storyboard,
                    "has_svg": has_svg,
                    "has_spec": has_spec,
                    "source_file": source_file,
                    "created_at": created_at,
                    "output_format": output_format,
                    "language": language,
                    "output_files": output_files,
                })
    return {"projects": projects}


@app.get("/api/projects/{project_name}/preview")
async def preview_project(project_name: str):
    """Preview storyboard content."""
    project_path = PROJECTS_DIR / project_name
    if not project_path.exists():
        raise HTTPException(404, "Project not found")
    
    storyboard_path = project_path / "storyboard.md"
    if not storyboard_path.exists():
        raise HTTPException(404, "No storyboard found")
    
    content = storyboard_path.read_text(encoding="utf-8")
    return {
        "name": project_name,
        "content": content,
        "source_file": (project_path / "source.txt").read_text(encoding="utf-8").strip() if (project_path / "source.txt").exists() else "",
        "created_at": (project_path / "created_at.txt").read_text(encoding="utf-8").strip() if (project_path / "created_at.txt").exists() else "",
    }


@app.get("/api/projects/{project_name}/download/{filename}")
async def download_project_file(project_name: str, filename: str):
    """Download a file from a project."""
    project_path = PROJECTS_DIR / project_name
    if not project_path.exists():
        raise HTTPException(404, "Project not found")
    
    file_path = project_path / filename
    if not file_path.exists():
        raise HTTPException(404, "File not found")
    
    return FileResponse(
        path=str(file_path),
        filename=filename,
    )


@app.delete("/api/projects/{project_name}")
async def delete_project(project_name: str):
    """Delete a project."""
    project_path = PROJECTS_DIR / project_name
    if not project_path.exists():
        raise HTTPException(404, "Project not found")
    shutil.rmtree(project_path)
    return {"status": "ok"}


# ─────────────────────────────────────────────
# LLM Status
# ─────────────────────────────────────────────

@app.get("/api/llm/status")
async def llm_status():
    """Check LLM availability."""
    settings = load_settings()
    llm_config = settings.get("llm", {})
    provider = llm_config.get("provider", "ollama")
    
    if provider == "ollama":
        endpoint = llm_config.get("endpoint", "http://localhost:11434")
        try:
            import httpx
            async with httpx.AsyncClient(timeout=5.0) as client:
                response = await client.get(f"{endpoint}/api/tags")
                response.raise_for_status()
                models = response.json().get("models", [])
                return {
                    "status": "connected",
                    "provider": "ollama",
                    "models": [m["name"] for m in models],
                }
        except Exception as e:
            return {
                "status": "disconnected",
                "provider": "ollama",
                "error": str(e),
            }
    
    return {"status": "configured", "provider": provider}


# ─────────────────────────────────────────────
# Pipeline — Auto PPTX Generation
# ─────────────────────────────────────────────

# In-memory pipeline status tracking
_pipeline_jobs: dict = {}


@app.post("/api/pipeline/run")
async def run_pipeline(
    file: UploadFile = File(...),
    style_prompt: str = Form(""),
    image_mode: str = Form("auto"),
    llm_provider: str = Form("openclaude"),
    tts_provider: str = Form("none"),
    tts_voice: str = Form(""),
    skip_audio: bool = Form(True),
):
    """Trigger the full PPTX generation pipeline."""
    import uuid
    job_id = str(uuid.uuid4())[:8]

    # Save uploaded file
    file_id = str(uuid.uuid4())[:8]
    upload_path = UPLOADS_DIR / f"{file_id}_{file.filename}"
    with open(upload_path, "wb") as f:
        content = await file.read()
        f.write(content)

    # Determine output dir
    original_name = Path(file.filename).stem
    output_dir = PROJECTS_DIR / f"{original_name}_{job_id}"
    output_dir.mkdir(exist_ok=True)

    _pipeline_jobs[job_id] = {
        "status": "running",
        "progress": "Starting pipeline...",
        "input_file": file.filename,
        "output_dir": str(output_dir),
        "result": None,
        "error": None,
    }

    # Run pipeline in background thread
    import threading

    def _run():
        try:
            # Add scripts dir to path
            scripts_dir = BASE_DIR / "skills" / "ppt-master" / "scripts"
            if str(scripts_dir) not in sys.path:
                sys.path.insert(0, str(scripts_dir))

            from auto_pptx_pipeline import run_pipeline as _run_pipeline

            _pipeline_jobs[job_id]["progress"] = "Extracting content..."

            result = _run_pipeline(
                input_file=str(upload_path),
                style_prompt=style_prompt,
                image_mode=image_mode,
                llm_provider=llm_provider,
                tts_provider=tts_provider if tts_provider != "none" else None,
                tts_voice=tts_voice or None,
                output_dir=str(output_dir),
                skip_audio=skip_audio,
            )

            _pipeline_jobs[job_id]["status"] = "completed"
            _pipeline_jobs[job_id]["result"] = result
            _pipeline_jobs[job_id]["progress"] = "Done!"

            # Find PPTX output
            pptx_files = list(output_dir.rglob("*.pptx"))
            if pptx_files:
                _pipeline_jobs[job_id]["pptx_path"] = str(pptx_files[0])

        except Exception as e:
            _pipeline_jobs[job_id]["status"] = "error"
            _pipeline_jobs[job_id]["error"] = str(e)
            _pipeline_jobs[job_id]["progress"] = f"Error: {str(e)[:200]}"

    thread = threading.Thread(target=_run, daemon=True)
    thread.start()

    return {"job_id": job_id, "status": "started", "output_dir": str(output_dir)}


@app.get("/api/pipeline/status/{job_id}")
async def pipeline_status(job_id: str):
    """Check pipeline job status."""
    if job_id not in _pipeline_jobs:
        raise HTTPException(404, "Job not found")
    return _pipeline_jobs[job_id]


@app.get("/api/pipeline/download/{job_id}")
async def pipeline_download(job_id: str):
    """Download the generated PPTX."""
    if job_id not in _pipeline_jobs:
        raise HTTPException(404, "Job not found")

    job = _pipeline_jobs[job_id]
    pptx_path = job.get("pptx_path")
    if not pptx_path or not Path(pptx_path).exists():
        raise HTTPException(404, "PPTX not ready yet")

    return FileResponse(
        path=pptx_path,
        filename=Path(pptx_path).name,
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
    )


@app.get("/api/pipeline/slides/{job_id}")
async def pipeline_slides(job_id: str):
    """Get slide data for preview."""
    if job_id not in _pipeline_jobs:
        raise HTTPException(404, "Job not found")

    job = _pipeline_jobs[job_id]
    output_dir = Path(job.get("output_dir", ""))

    enriched_path = output_dir / "slides_enriched.json"
    if not enriched_path.exists():
        return {"slides": [], "status": job["status"]}

    slides = json.loads(enriched_path.read_text(encoding="utf-8"))

    # Add image paths
    for slide in slides:
        num = slide.get("slide_number", 0)
        slide_dir = output_dir / "slides" / f"slide_{num:02d}"
        images = list((slide_dir / "images").glob("*")) if (slide_dir / "images").exists() else []
        slide["local_images"] = [str(img) for img in images]

    return {"slides": slides, "status": job["status"], "total": len(slides)}


@app.get("/api/pipeline/slide-image/{job_id}/{slide_num}")
async def pipeline_slide_image(job_id: str, slide_num: int):
    """Get slide image for preview."""
    if job_id not in _pipeline_jobs:
        raise HTTPException(404, "Job not found")

    job = _pipeline_jobs[job_id]
    output_dir = Path(job.get("output_dir", ""))
    images_dir = output_dir / "slides" / f"slide_{slide_num:02d}" / "images"

    if not images_dir.exists():
        raise HTTPException(404, "No images for this slide")

    # Only pick actual image files, not directories
    image_extensions = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp"}
    images = [f for f in images_dir.iterdir() if f.is_file() and f.suffix.lower() in image_extensions]
    if not images:
        raise HTTPException(404, "No images found")

    return FileResponse(path=str(images[0]))


@app.get("/api/pipeline/download-folder/{job_id}")
async def pipeline_download_folder(job_id: str):
    """Download the slides folder as a ZIP file."""
    import zipfile
    import io

    if job_id not in _pipeline_jobs:
        raise HTTPException(404, "Job not found")

    job = _pipeline_jobs[job_id]
    output_dir = Path(job.get("output_dir", ""))
    slides_dir = output_dir / "slides"

    if not slides_dir.exists():
        raise HTTPException(404, "Slides folder not ready")

    # Create ZIP in memory
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        # Add slides folder
        for slide_dir in sorted(slides_dir.iterdir()):
            if slide_dir.is_dir():
                for file in slide_dir.rglob("*"):
                    if file.is_file():
                        arcname = f"slides/{slide_dir.name}/{file.name}"
                        zf.write(file, arcname)

        # Add PPTX if exists
        pptx_path = job.get("pptx_path")
        if pptx_path and Path(pptx_path).exists():
            zf.write(pptx_path, f"exports/{Path(pptx_path).name}")

        # Add enriched content
        enriched_path = output_dir / "slides_enriched.json"
        if enriched_path.exists():
            zf.write(enriched_path, "slides_enriched.json")

        # Add audio folder
        audio_dir = output_dir / "audio"
        if audio_dir.exists():
            for audio_file in audio_dir.glob("*.mp3"):
                zf.write(audio_file, f"audio/{audio_file.name}")

    zip_buffer.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f"attachment; filename=slides_{job_id}.zip"}
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
